from __future__ import annotations

import base64
import tempfile
import unittest
from pathlib import Path

from docx import Document

from a3presentation.domain.api import ChartOverride
from a3presentation.domain.chart import ChartConfidence, ChartSeries, ChartSpec, ChartType
from a3presentation.domain.presentation import PresentationPlan, SlideKind, SlideSpec
from a3presentation.domain.template import GenerationMode, PlaceholderKind
from a3presentation.services.deck_audit import audit_generated_presentation, find_capacity_violations
from a3presentation.services.document_text_extractor import DocumentTextExtractor
from a3presentation.services.planner import TextToPlanService
from a3presentation.services.pptx_generator import PptxGenerator
from a3presentation.services.template_analyzer import TemplateAnalyzer
from a3presentation.services.template_registry import TemplateRegistry
from a3presentation.settings import get_settings


SMALL_PNG_BYTES = base64.b64decode(
    "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO9W6i8AAAAASUVORK5CYII="
)


class QualityContractTests(unittest.TestCase):
    @classmethod
    def setUpClass(cls) -> None:
        settings = get_settings()
        registry = TemplateRegistry(settings.templates_dir)
        cls.manifest = registry.get_template("corp_light_v1")
        cls.template_path = registry.get_template_pptx_path("corp_light_v1")
        cls.extractor = DocumentTextExtractor()
        cls.planner = TextToPlanService()
        cls.generator = PptxGenerator()
        cls.registry = registry
        cls.settings = settings
        cls.analyzer = TemplateAnalyzer()

    def test_text_only_document_respects_quality_contract(self) -> None:
        raw_text = (
            "# Стратегия\n\n"
            "## Контекст\n"
            "Компания усиливает платформу и перераспределяет ресурсы в направления с лучшей окупаемостью.\n\n"
            "## Решения\n"
            "- Ускорить развитие core-платформы\n"
            "- Усилить партнёрские интеграции\n"
            "- Повысить долю recurring revenue\n"
        )
        text, tables, blocks = self.extractor.extract("strategy.md", raw_text.encode("utf-8"))
        plan = self.planner.build_plan("corp_light_v1", text, None, tables, blocks)

        violations = self._generate_and_audit(plan)
        self.assertEqual(violations, [])

    def test_chart_heavy_document_respects_quality_contract(self) -> None:
        document = Document()
        document.add_paragraph("Отчёт по каналам продаж")
        document.add_heading("1. Итоги квартала", level=1)
        document.add_paragraph("Основной рост пришёл из платных каналов и рекомендаций.")
        table = document.add_table(rows=4, cols=2)
        table.cell(0, 0).text = "Канал"
        table.cell(0, 1).text = "Выручка"
        table.cell(1, 0).text = "SEO"
        table.cell(1, 1).text = "120"
        table.cell(2, 0).text = "Ads"
        table.cell(2, 1).text = "200"
        table.cell(3, 0).text = "Referral"
        table.cell(3, 1).text = "90"
        document.add_heading("2. Вывод", level=1)
        document.add_paragraph("Нужно перераспределить бюджет в пользу каналов с лучшей окупаемостью.")

        with tempfile.TemporaryDirectory() as temp_dir:
            docx_path = Path(temp_dir) / "chart-heavy.docx"
            document.save(docx_path)
            text, tables, blocks = self.extractor.extract(docx_path.name, docx_path.read_bytes())

        plan = self.planner.build_plan(
            "corp_light_v1",
            text,
            None,
            tables,
            blocks,
            chart_overrides=[
                ChartOverride(
                    table_id="table_1",
                    mode="chart",
                    selected_chart=ChartSpec(
                        chart_id="chart_1",
                        source_table_id="table_1",
                        chart_type=ChartType.COLUMN,
                        title="Выручка по каналам",
                        categories=["SEO", "Ads", "Referral"],
                        series=[ChartSeries(name="Выручка", values=[120.0, 200.0, 90.0])],
                        confidence=ChartConfidence.HIGH,
                    ),
                )
            ],
        )

        violations = self._generate_and_audit(plan)
        self.assertEqual(violations, [])

    def test_image_heavy_document_respects_quality_contract(self) -> None:
        image_b64 = base64.b64encode(SMALL_PNG_BYTES).decode("ascii")
        blocks = [
            self._block("paragraph", "Отчет по продукту"),
            self._block("heading", "1. Схема процесса", level=1),
            self._block("paragraph", "Ниже приведена ключевая схема целевого процесса."),
            self._block(
                "image",
                "Схема целевого процесса",
                image_name="process.png",
                image_content_type="image/png",
                image_base64=image_b64,
            ),
            self._block("heading", "2. Выводы", level=1),
            self._block("paragraph", "Иллюстрация подтверждает узкие места и точки автоматизации."),
        ]
        raw_text = "\n".join(block.text or "" for block in blocks)
        plan = self.planner.build_plan("corp_light_v1", raw_text, None, [], blocks)

        violations = self._generate_and_audit(plan)
        self.assertEqual(violations, [])

    def test_uploaded_layout_template_text_document_respects_quality_contract(self) -> None:
        template_id = "razmeshchenie_soglasiy"
        template_path = self.settings.templates_dir / template_id / "template.pptx"
        if not template_path.exists():
            self.skipTest(f"optional uploaded template is not installed: {template_id}")
        manifest = self.analyzer.analyze(
            template_id=template_id,
            template_path=template_path,
            display_name="Размещение согласий",
        )
        manifest.generation_mode = GenerationMode.LAYOUT
        layout = next(item for item in manifest.layouts if item.key == "титульный_слайд")
        plan = PresentationPlan(
            template_id=template_id,
            title="Uploaded Layout Quality",
            slides=[
                SlideSpec(
                    kind=SlideKind.TEXT,
                    title="Контекст",
                    text="Analyzer-derived layout metadata должна проходить quality gate на свежем артефакте.",
                    preferred_layout_key=layout.key,
                )
            ],
        )

        violations = self._generate_and_audit(plan, manifest=manifest, template_path=template_path)
        self.assertEqual(violations, [])

    def test_uploaded_prototype_template_text_document_respects_quality_contract(self) -> None:
        template_id = "razmeshchenie_soglasiy"
        if not (self.settings.templates_dir / template_id).exists():
            self.skipTest(f"optional uploaded template is not installed: {template_id}")
        manifest = self.registry.get_template(template_id)
        template_path = self.settings.templates_dir / template_id / manifest.source_pptx
        prototype = next(
            item for item in manifest.prototype_slides if any(token.binding == "main_text" for token in item.tokens)
        )
        plan = PresentationPlan(
            template_id=template_id,
            title="Uploaded Prototype Quality",
            slides=[
                SlideSpec(
                    kind=SlideKind.TEXT,
                    title="Основные выводы",
                    text=(
                        "Template-aware quality gate должна читать геометрию и margins из prototype token metadata "
                        "и подтверждать это на свежей генерации пользовательского шаблона."
                    ),
                    notes="Служебная строка для footer prototype.",
                    preferred_layout_key=prototype.key,
                )
            ],
        )

        violations = self._generate_and_audit(plan, manifest=manifest, template_path=template_path)
        self.assertEqual(violations, [])

    def test_uploaded_prototype_chart_template_respects_quality_contract(self) -> None:
        manifests = [
            manifest
            for manifest in self.registry.list_templates()
            if manifest.generation_mode == GenerationMode.PROTOTYPE
            and any(any(token.binding == "chart_image" for token in prototype.tokens) for prototype in manifest.prototype_slides)
        ]
        if not manifests:
            self.skipTest("no optional uploaded prototype template with chart_image binding is installed")

        for manifest in manifests:
            template_path = self.settings.templates_dir / manifest.template_id / manifest.source_pptx
            if not template_path.exists():
                continue

            with self.subTest(template_id=manifest.template_id):
                plan = PresentationPlan(
                    template_id=manifest.template_id,
                    title=f"{manifest.display_name} Prototype Chart Quality",
                    slides=[
                        SlideSpec(
                            kind=SlideKind.TITLE,
                            title=f"{manifest.display_name} Prototype Chart Quality",
                            preferred_layout_key="cover",
                        ),
                        SlideSpec(
                            kind=SlideKind.CHART,
                            title="Выручка и маржа",
                            subtitle="Template-aware chart quality gate для uploaded prototype template",
                            chart=ChartSpec(
                                chart_id="quality_chart_prototype",
                                source_table_id="table_1",
                                chart_type=ChartType.COMBO,
                                title="Выручка и маржа",
                                categories=["Q1", "Q2", "Q3"],
                                series=[
                                    ChartSeries(name="Выручка", values=[104_300_000.0, 111_300_000.0, 135_700_000.0], unit="RUB"),
                                    ChartSeries(name="Маржа", values=[18.0, 22.0, 27.0], unit="%"),
                                ],
                                confidence=ChartConfidence.HIGH,
                                value_format="number",
                            ),
                            preferred_layout_key="table",
                        ),
                    ],
                )

                violations = self._generate_and_audit(plan, manifest=manifest, template_path=template_path)
                self.assertEqual(violations, [])

    def test_installed_user_templates_keep_optional_quality_matrix_without_fatal_slot_regressions(self) -> None:
        manifests = [manifest for manifest in self.registry.list_templates() if manifest.template_id != "corp_light_v1"]
        if not manifests:
            self.skipTest("no optional user templates are installed")

        checked_templates = 0
        for manifest in manifests:
            template_path = self.settings.templates_dir / manifest.template_id / manifest.source_pptx
            if not template_path.exists():
                continue

            text_layout_key = self._first_supported_template_key(manifest, {SlideKind.TEXT, SlideKind.BULLETS, SlideKind.TWO_COLUMN})
            if text_layout_key is None:
                continue

            with self.subTest(template_id=manifest.template_id):
                slides = [
                    SlideSpec(
                        kind=SlideKind.TEXT,
                        title="Контекст",
                        text="Короткий контрольный текст для template-aware quality sweep.",
                        preferred_layout_key=text_layout_key,
                    )
                ]
                chart_layout_key = self._first_supported_template_key(manifest, {SlideKind.CHART})
                if chart_layout_key is not None:
                    slides.append(
                        SlideSpec(
                            kind=SlideKind.CHART,
                            title="Выручка",
                            subtitle="Контрольный chart sweep",
                            chart=ChartSpec(
                                chart_id=f"quality_matrix_{manifest.template_id}",
                                source_table_id="table_1",
                                chart_type=ChartType.COMBO,
                                title="Выручка и маржа",
                                categories=["Q1", "Q2", "Q3"],
                                series=[
                                    ChartSeries(name="Выручка", values=[104_300_000.0, 111_300_000.0, 135_700_000.0], unit="RUB"),
                                    ChartSeries(name="Маржа", values=[18.0, 22.0, 27.0], unit="%"),
                                ],
                                confidence=ChartConfidence.HIGH,
                                value_format="number",
                            ),
                            preferred_layout_key=chart_layout_key,
                        )
                    )

                plan = PresentationPlan(
                    template_id=manifest.template_id,
                    title=f"{manifest.display_name} Optional Quality Matrix",
                    slides=slides,
                )
                violations = self._generate_and_audit(plan, manifest=manifest, template_path=template_path)
                forbidden_rules = {"missing_table_shape", "missing_chart_shape", "missing_image_shape", "content_order_mismatch"}
                self.assertFalse(any(item.rule in forbidden_rules for item in violations))
                checked_templates += 1

        self.assertGreater(checked_templates, 0)

    def _generate_and_audit(self, plan, manifest=None, template_path=None):
        active_manifest = manifest or self.manifest
        active_template_path = template_path or self.template_path
        with tempfile.TemporaryDirectory() as temp_dir:
            output_path = self.generator.generate(
                template_path=active_template_path,
                manifest=active_manifest,
                plan=plan,
                output_dir=Path(temp_dir),
            )
            audits = audit_generated_presentation(output_path, plan, active_manifest)
        return find_capacity_violations(audits)

    def _first_supported_template_key(self, manifest, slide_kinds: set[SlideKind]) -> str | None:
        supported_kinds = {kind.value for kind in slide_kinds}
        wants_chart = SlideKind.CHART in slide_kinds
        if manifest.generation_mode == GenerationMode.PROTOTYPE:
            if wants_chart:
                for prototype in manifest.prototype_slides:
                    if "chart" in prototype.supported_slide_kinds and any(token.binding == "chart_image" for token in prototype.tokens):
                        return prototype.key

            preferred_text_bindings = ("main_text", "body", "bullets", "left_text", "left_bullets")
            for binding in preferred_text_bindings:
                for prototype in manifest.prototype_slides:
                    if supported_kinds & set(prototype.supported_slide_kinds) and any(token.binding == binding for token in prototype.tokens):
                        return prototype.key

            for prototype in manifest.prototype_slides:
                if supported_kinds & set(prototype.supported_slide_kinds):
                    return prototype.key
            return None

        if wants_chart:
            for layout in manifest.layouts:
                if "chart" in layout.supported_slide_kinds and any(
                    placeholder.kind in {PlaceholderKind.CHART, PlaceholderKind.IMAGE} for placeholder in layout.placeholders
                ):
                    return layout.key

        for layout in manifest.layouts:
            if supported_kinds & set(layout.supported_slide_kinds) and any(
                placeholder.kind == PlaceholderKind.BODY for placeholder in layout.placeholders
            ):
                return layout.key

        for layout in manifest.layouts:
            if supported_kinds & set(layout.supported_slide_kinds):
                return layout.key
        return None

    def _block(self, kind: str, text: str, **kwargs):
        from a3presentation.domain.api import DocumentBlock

        return DocumentBlock(kind=kind, text=text, **kwargs)


if __name__ == "__main__":
    unittest.main()
