from __future__ import annotations

from io import BytesIO
import unittest

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

from a3presentation.services.document_text_extractor import DocumentTextExtractor


class DocumentTextExtractorTests(unittest.TestCase):
    def test_markdown_extracts_headings_and_lists(self) -> None:
        extractor = DocumentTextExtractor()
        markdown = (
            "# Рынок медицины для A3\n\n"
            "## Основные сегменты\n"
            "- Частные клиники\n"
            "- Телемедицина\n\n"
            "Рынок быстро растет за счет цифровизации и частного спроса.\n"
        ).encode("utf-8")

        text, tables, blocks = extractor.extract("sample.md", markdown)

        self.assertTrue(text)
        self.assertEqual(tables, [])
        self.assertGreaterEqual(len(blocks), 3)
        self.assertEqual(blocks[0].kind, "heading")
        self.assertEqual(blocks[0].text, "Рынок медицины для A3")
        self.assertEqual(blocks[1].kind, "subheading")
        self.assertEqual(blocks[2].kind, "list")
        self.assertEqual(blocks[2].items, ["Частные клиники", "Телемедицина"])

    def test_markdown_numbered_list_is_not_misclassified_as_headings(self) -> None:
        extractor = DocumentTextExtractor()
        markdown = (
            "# План\n\n"
            "1. Первый шаг\n"
            "2. Второй шаг\n"
        ).encode("utf-8")

        _, _, blocks = extractor.extract("sample.md", markdown)

        self.assertEqual(blocks[0].kind, "heading")
        self.assertEqual(blocks[1].kind, "list")
        self.assertEqual(blocks[1].items, ["Первый шаг", "Второй шаг"])

    def test_docx_detects_outline_and_quarter_headings_without_explicit_heading_style(self) -> None:
        extractor = DocumentTextExtractor()
        document = Document()
        document.add_paragraph("A3")
        document.add_paragraph("1.4 Показатели 2025")
        document.add_paragraph("Ключевые финансовые результаты.")
        document.add_paragraph("Q2 2026")
        document.add_paragraph("Переход к следующему этапу.")

        buffer = BytesIO()
        document.save(buffer)

        _, _, blocks = extractor.extract("sample.docx", buffer.getvalue())

        self.assertEqual(blocks[1].kind, "subheading")
        self.assertEqual(blocks[1].level, 2)
        self.assertEqual(blocks[3].kind, "subheading")
        self.assertEqual(blocks[3].level, 2)

    def test_docx_merges_short_narrative_paragraph_continuations_into_one_block(self) -> None:
        extractor = DocumentTextExtractor()
        document = Document()
        document.add_heading("1. Vision и стратегические цели", level=1)
        document.add_heading("1.1 Vision 2030", level=2)
        document.add_paragraph("Стать одной из ведущих продуктовых финтех IT-компаний в России:")
        document.add_paragraph("устойчивой, быстрорастущей, узнаваемой и диверсифицированной")
        document.add_paragraph("по продуктам, партнёрам и источникам выручки.")

        buffer = BytesIO()
        document.save(buffer)

        _, _, blocks = extractor.extract("sample.docx", buffer.getvalue())

        narrative_blocks = [block for block in blocks if block.kind == "paragraph"]
        self.assertEqual(len(narrative_blocks), 1)
        self.assertEqual(
            narrative_blocks[0].text,
            (
                "Стать одной из ведущих продуктовых финтех IT-компаний в России: "
                "устойчивой, быстрорастущей, узнаваемой и диверсифицированной "
                "по продуктам, партнёрам и источникам выручки."
            ),
        )

    def test_docx_does_not_merge_independent_short_paragraph_sentences(self) -> None:
        extractor = DocumentTextExtractor()
        document = Document()
        document.add_heading("2. Роль бренда", level=1)
        document.add_paragraph("Инфраструктурный игрок, объединяющий участников финансового рынка.")
        document.add_paragraph("Бренд должен быть узнаваемым и технологичным.")

        buffer = BytesIO()
        document.save(buffer)

        _, _, blocks = extractor.extract("sample.docx", buffer.getvalue())

        narrative_blocks = [block for block in blocks if block.kind == "paragraph"]
        self.assertEqual(len(narrative_blocks), 2)

    def test_docx_does_not_merge_reference_tail_into_narrative_paragraph(self) -> None:
        extractor = DocumentTextExtractor()
        document = Document()
        document.add_heading("1. Контекст", level=1)
        document.add_paragraph("Основной текст раздела должен остаться narrative-блоком")
        document.add_paragraph("[1] https://example.com/source-1")

        buffer = BytesIO()
        document.save(buffer)

        _, _, blocks = extractor.extract("sample.docx", buffer.getvalue())

        narrative_blocks = [block for block in blocks if block.kind == "paragraph"]
        self.assertEqual(len(narrative_blocks), 2)
        self.assertEqual(narrative_blocks[0].text, "Основной текст раздела должен остаться narrative-блоком")
        self.assertEqual(narrative_blocks[1].text, "[1] https://example.com/source-1")

    def test_docx_extracts_table_cell_fill_colors(self) -> None:
        extractor = DocumentTextExtractor()
        document = Document()
        table = document.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Показатель"
        table.cell(0, 1).text = "Значение"
        table.cell(1, 0).text = "Выручка"
        table.cell(1, 1).text = "120"
        table.cell(2, 0).text = "Маржа"
        table.cell(2, 1).text = "24%"
        self._set_cell_fill(table.cell(0, 0), "1F4E79")
        self._set_cell_fill(table.cell(0, 1), "1F4E79")
        self._set_cell_fill(table.cell(1, 1), "D9EAF7")

        buffer = BytesIO()
        document.save(buffer)

        _, tables, blocks = extractor.extract("sample.docx", buffer.getvalue())

        self.assertEqual(len(tables), 1)
        self.assertEqual(tables[0].header_fill_colors, ["1F4E79", "1F4E79"])
        self.assertEqual(tables[0].row_fill_colors[0], [None, "D9EAF7"])
        table_blocks = [block for block in blocks if block.kind == "table"]
        self.assertEqual(table_blocks[0].table.header_fill_colors, ["1F4E79", "1F4E79"])

    def _set_cell_fill(self, cell, color: str) -> None:
        tc_pr = cell._tc.get_or_add_tcPr()
        shading = tc_pr.find(qn("w:shd"))
        if shading is None:
            shading = OxmlElement("w:shd")
            tc_pr.append(shading)
        shading.set(qn("w:fill"), color)


if __name__ == "__main__":
    unittest.main()
